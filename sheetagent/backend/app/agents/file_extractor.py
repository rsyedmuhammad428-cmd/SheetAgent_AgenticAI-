"""
SheetAgent AI — file_extractor.py (v6 FINAL)

KEY FIXES:

1. MULTI-PAGE PDF (main bug from screenshots):
   - Smart repeat-header detection: if page 2's table[0] matches the column
     headers from page 1, it is skipped (it's just a repeated header row).
     If it does NOT match, it is treated as a data row and included.
   - After pdfplumber table extraction, if we also collected text from all pages,
     we run Gemini text extraction on the full text and MERGE any rows that
     table extraction missed. This catches highlighted/coloured rows that
     pdfplumber may have missed as separate table objects.

2. IMAGE EXTRACTION (PNG/JPG/JPEG kept — improved pipeline):
   - 3-pass: JSON Vision → OCR Vision → Gemini text parse on OCR result
   - Added pytesseract as an optional 4th fallback (if installed).
   - Images are kept in the supported formats — the Vision pipeline is now
     robust enough to handle data tables in images.

3. ALL PAGES: pdfplumber reads up to 100 pages, pdf2image renders up to 20.
"""
import asyncio
import json
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


class FileExtractor:

    async def extract(self, file_path: str, session_id: str = "") -> dict:
        fp  = Path(file_path)
        ext = fp.suffix.lower()

        result = {
            "file_name":         fp.name,
            "file_type":         ext,
            "extracted_text":    "",
            "extracted_data":    [],
            "extracted_columns": [],
            "row_count":         0,
            "file_summary":      "",
            "success":           False,
        }

        if not fp.exists():
            return self._empty_result(f"File not found: {file_path}")

        try:
            if ext == ".csv":
                await self._extract_csv(fp, result)
            elif ext in (".xlsx", ".xls", ".xlsm"):
                await self._extract_excel(fp, result)
            elif ext == ".pdf":
                await self._extract_pdf(fp, result, session_id)
            elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".webp", ".bmp"):
                await self._extract_image(fp, result, session_id)
            elif ext in (".docx", ".doc"):
                await self._extract_word(fp, result)
            elif ext == ".txt":
                await self._extract_text(fp, result)
            else:
                result["extracted_text"] = f"Unsupported file type: {ext}"
                result["file_summary"]   = f"File: {fp.name} (unsupported type)"
                return result

            result["success"]   = True
            result["row_count"] = len(result["extracted_data"])
            result["file_summary"] = self._build_summary(result)

        except Exception as e:
            logger.error(f"[Extractor] Error on {fp.name}: {e}", exc_info=True)
            result["extracted_text"] = f"Error reading {fp.name}: {str(e)}"
            result["file_summary"]   = f"Could not read {fp.name}: {str(e)}"
            raise

        return result

    # ── CSV ───────────────────────────────────────────────────────────────────

    async def _extract_csv(self, fp: Path, result: dict):
        import pandas as pd
        df = pd.read_csv(fp, encoding="utf-8", on_bad_lines="skip")
        df = df.where(df.notna(), None)
        result["extracted_data"]    = df.to_dict(orient="records")
        result["extracted_columns"] = list(df.columns)
        result["extracted_text"]    = (
            f"CSV — {len(df)} rows, {len(df.columns)} columns.\n"
            f"Columns: {list(df.columns)}\nRows: {len(df)}\n"
            f"Sample (first 10):\n{df.head(10).to_string()}"
        )

    # ── Excel ─────────────────────────────────────────────────────────────────

    async def _extract_excel(self, fp: Path, result: dict):
        import pandas as pd
        xl       = pd.ExcelFile(fp, engine="openpyxl")
        all_data, all_cols, parts = [], [], []
        for sheet_name in xl.sheet_names[:5]:
            df = pd.read_excel(fp, sheet_name=sheet_name, engine="openpyxl")
            df = df.where(df.notna(), None)
            all_data.extend(df.to_dict(orient="records"))
            if not all_cols:
                all_cols = list(df.columns)
            parts.append(
                f"Sheet '{sheet_name}': {len(df)} rows, cols: {list(df.columns)}\n"
                f"{df.head(10).to_string()}"
            )
        result["extracted_data"]    = all_data
        result["extracted_columns"] = all_cols
        result["extracted_text"]    = "\n\n".join(parts)

    # ── PDF ───────────────────────────────────────────────────────────────────

    async def _extract_pdf(self, fp: Path, result: dict, session_id: str):
        """
        3-stage PDF extraction with FIXED multi-page support:

        Stage 1: pdfplumber table extraction — reads ALL pages.
                 Smart header detection: repeated header rows on subsequent pages
                 are skipped; first row of page 2+ is included as DATA if it
                 doesn't match the column headers.

        Stage 2: pdfplumber text + Gemini text parse — for text PDFs with no
                 table objects, AND as a supplement to Stage 1 to catch any
                 rows missed across page boundaries.

        Stage 3: pdf2image + Vision — last resort for scanned PDFs.
        """
        try:
            import pdfplumber
            texts, all_data, all_cols = [], [], []
            known_headers_set = set()   # lowercased header values for comparison

            with pdfplumber.open(str(fp)) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"[Extractor] PDF has {total_pages} pages")

                for page_idx, page in enumerate(pdf.pages[:100]):
                    # ── Collect text from every page ──────────────────────
                    text = page.extract_text() or ""
                    if text.strip():
                        texts.append(f"[Page {page_idx + 1}]\n{text.strip()}")

                    # ── Extract tables from every page ────────────────────
                    for table in page.extract_tables():
                        if not table or not table[0]:
                            continue

                        raw_first_row = [str(c or "").strip() for c in table[0]]

                        # On page 1: first row is the header
                        if page_idx == 0 and not all_cols:
                            all_cols = raw_first_row
                            known_headers_set = {h.lower() for h in all_cols if h}
                            data_rows = table[1:]
                            logger.info(f"[Extractor] Page 1 headers: {all_cols}")

                        # On subsequent pages: check if first row is a repeat header
                        elif page_idx > 0 or all_cols:
                            first_row_lower = {c.lower() for c in raw_first_row if c}
                            overlap = first_row_lower & known_headers_set
                            is_repeat_header = (
                                len(overlap) >= max(1, len(known_headers_set) // 2)
                            )
                            if is_repeat_header:
                                # Skip repeated header, take data from row 1 onward
                                data_rows = table[1:]
                                logger.info(
                                    f"[Extractor] Page {page_idx+1}: repeated header "
                                    f"skipped, {len(data_rows)} data rows"
                                )
                            else:
                                # First row is actual data (no header on this page)
                                data_rows = table
                                logger.info(
                                    f"[Extractor] Page {page_idx+1}: no header row, "
                                    f"all {len(data_rows)} rows are data"
                                )
                        else:
                            data_rows = table[1:]

                        # Add data rows
                        for row in data_rows:
                            if not row:
                                continue
                            if not all_cols:
                                continue
                            row_dict = {
                                all_cols[j]: _coerce(str(cell or "").strip())
                                for j, cell in enumerate(row)
                                if j < len(all_cols)
                            }
                            if any(v != "" for v in row_dict.values()):
                                all_data.append(row_dict)

            logger.info(
                f"[Extractor] pdfplumber: {len(all_data)} rows from "
                f"{len(texts)} pages of text"
            )

            # ── Stage 1 success: got table data ───────────────────────────
            if all_data:
                full_text = "\n\n".join(texts)
                result["extracted_data"]    = all_data
                result["extracted_columns"] = all_cols
                result["extracted_text"]    = full_text or str(all_data[:3])

                # ── Stage 1 supplement: if we also have text, run Gemini ──
                # on it and merge any rows that table extraction missed
                # (handles highlighted/coloured rows that pdfplumber misses)
                if texts and len(all_data) < 200:
                    logger.info("[Extractor] Supplementing table data with Gemini text parse")
                    try:
                        gemini_rows = await _gemini_extract_structured(full_text)
                        if gemini_rows and len(gemini_rows) > len(all_data):
                            merged = _merge_rows(all_data, gemini_rows, all_cols)
                            if len(merged) > len(all_data):
                                logger.info(
                                    f"[Extractor] Merged: {len(all_data)} table rows + "
                                    f"{len(gemini_rows)} Gemini rows = {len(merged)} total"
                                )
                                all_data = merged
                                result["extracted_data"] = all_data
                                result["row_count"]      = len(all_data)
                    except Exception as e:
                        from app.agents.quota_helper import is_quota_error
                        if is_quota_error(e):
                            raise
                        logger.warning(f"[Extractor] Gemini supplement failed: {e}")

                return

            # ── Stage 2: text-only PDF ─────────────────────────────────────
            if texts:
                full_text = "\n\n".join(texts)
                result["extracted_text"] = full_text

                # 2a: fast regex parse
                parsed = _parse_text_as_table(full_text)
                if parsed["rows"]:
                    result["extracted_data"]    = parsed["rows"]
                    result["extracted_columns"] = parsed["headers"]
                    logger.info(f"[Extractor] PDF regex: {len(parsed['rows'])} rows")
                    return

                # 2b: Gemini text extraction
                logger.info("[Extractor] PDF text — Gemini text parse")
                gemini_rows = await _gemini_extract_structured(full_text)
                if gemini_rows:
                    result["extracted_data"]    = gemini_rows
                    result["extracted_columns"] = list(gemini_rows[0].keys())
                    logger.info(f"[Extractor] PDF Gemini text: {len(gemini_rows)} rows")
                    return

                logger.info("[Extractor] PDF text extracted but no rows")
                return

        except ImportError:
            logger.warning("[Extractor] pdfplumber not installed — using Vision")
        except Exception as e:
            from app.agents.quota_helper import is_quota_error
            if is_quota_error(e):
                raise
            logger.warning(f"[Extractor] pdfplumber error: {e} — using Vision")

        # Stage 3: scanned PDF → Vision
        logger.info("[Extractor] PDF no text — Vision on page images")
        await self._extract_pdf_via_vision(fp, result, session_id)

    async def _extract_pdf_via_vision(self, fp: Path, result: dict, session_id: str):
        try:
            import tempfile, os
            from pdf2image import convert_from_path
            images = convert_from_path(str(fp), first_page=1, last_page=20, dpi=150)
            if not images:
                result["extracted_text"] = "Empty PDF"
                return

            all_data, all_cols, all_texts = [], [], []
            known_headers_set = set()

            for i, img in enumerate(images):
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    img.save(tmp.name, "PNG")
                    tmp_path = tmp.name

                page_result = {
                    "extracted_text": "", "extracted_data": [], "extracted_columns": []
                }
                await self._extract_image(Path(tmp_path), page_result, session_id)
                os.unlink(tmp_path)

                if page_result["extracted_data"]:
                    page_rows  = page_result["extracted_data"]
                    page_cols  = page_result["extracted_columns"]

                    if not all_cols:
                        all_cols = page_cols
                        known_headers_set = {c.lower() for c in all_cols if c}
                        all_data.extend(page_rows)
                    else:
                        # Check if first row of this page is a repeated header
                        if page_rows:
                            first_vals = {str(v).lower() for v in page_rows[0].values() if v}
                            is_repeat  = len(first_vals & known_headers_set) >= max(1, len(known_headers_set) // 2)
                            rows_to_add = page_rows[1:] if is_repeat else page_rows
                            all_data.extend(rows_to_add)

                if page_result["extracted_text"]:
                    all_texts.append(f"[Page {i+1}]\n{page_result['extracted_text']}")

            result["extracted_data"]    = all_data
            result["extracted_columns"] = all_cols
            result["extracted_text"]    = "\n\n".join(all_texts)
            logger.info(f"[Extractor] PDF Vision: {len(all_data)} rows from {len(images)} pages")

        except ImportError:
            result["extracted_text"] = (
                "Scanned PDF detected but pdf2image is not installed.\n"
                "Install it with: pip install pdf2image"
            )
        except Exception as e:
            from app.agents.quota_helper import is_quota_error
            if is_quota_error(e):
                raise
            logger.error(f"[Extractor] PDF Vision failed: {e}")
            result["extracted_text"] = f"PDF Vision extraction failed: {e}"

    # ── Image (PNG/JPG/JPEG/etc) ──────────────────────────────────────────────

    async def _extract_image(self, fp: Path, result: dict, session_id: str):
        """
        3-pass + optional 4th pass (pytesseract):

        Pass 1 — JSON Vision prompt: ask Gemini to return structured JSON.
        Pass 2 — OCR Vision: verbatim text transcription → 5-layer parser.
        Pass 3 — Gemini text extraction on OCR output.
        Pass 4 — pytesseract local OCR (if installed, no API call needed).
        """
        try:
            import google.generativeai as genai
            import PIL.Image
            from app.config import settings
            from app.agents.rate_limiter import _vision_call_with_retry

            genai.configure(api_key=settings.gemini_api_key)
            model = genai.GenerativeModel("gemini-1.5-flash")
            img   = PIL.Image.open(str(fp))

            # ── Pass 1: Structured JSON ───────────────────────────────────
            json_prompt = """You are a data extraction expert. Look at this image carefully.

This image may show a spreadsheet, a data table, or text containing structured data.

YOUR TASK: Extract every row of data and return as JSON.

STRICT RULES:
1. Copy EVERY value EXACTLY as it appears. Never invent or replace any value.
2. Do NOT use Alice, Bob, Item A, Field 1, or any placeholder. Use actual values.
3. Use the EXACT column header text from the image.
4. Extract ALL rows visible — never skip or truncate.
5. Numbers must be actual numbers: 125.50 not "$125.50".

Return ONLY this JSON (no markdown, no code fences, start with {):
{
  "headers": ["ExactHeader1", "ExactHeader2"],
  "rows": [["actualValue", 123], ["actualValue", 456]],
  "summary": "What this image shows"
}"""

            response1 = await _vision_call_with_retry(model, [img, json_prompt])
            raw1 = response1.text.strip() if response1.text else ""
            logger.info(f"[Extractor] Vision P1 ({len(raw1)} chars): {raw1[:200]}")

            rows, headers, summary = _parse_vision_response(raw1)
            logger.info(f"[Extractor] P1 parsed: {len(rows)} rows")

            # ── Pass 2: Verbatim OCR ─────────────────────────────────────
            ocr_text = ""
            if not rows:
                ocr_prompt = (
                    "Read every character in this image exactly as written. "
                    "Output all text, numbers, symbols, and punctuation. "
                    "Preserve line breaks. Do not summarize anything."
                )
                response2 = await _vision_call_with_retry(model, [img, ocr_prompt])
                ocr_text  = response2.text.strip() if response2.text else ""
                logger.info(f"[Extractor] P2 OCR ({len(ocr_text)} chars): {ocr_text[:200]}")

                if ocr_text:
                    rows, headers, summary = _parse_vision_response(ocr_text)
                    logger.info(f"[Extractor] P2 parsed: {len(rows)} rows")

            # ── Pass 3: Gemini text extraction on OCR ────────────────────
            if not rows and ocr_text:
                logger.info("[Extractor] P3: Gemini text extraction on OCR")
                rows = await _gemini_extract_structured(ocr_text)
                if rows:
                    headers = list(rows[0].keys())
                    summary = "Extracted via OCR + Gemini"
                    logger.info(f"[Extractor] P3: {len(rows)} rows")

            # ── Pass 4: pytesseract local OCR (no API call) ───────────────
            if not rows:
                logger.info("[Extractor] P4: trying pytesseract")
                try:
                    import pytesseract
                    tess_text = pytesseract.image_to_string(img)
                    if tess_text.strip():
                        rows_t, headers_t, _ = _parse_vision_response(tess_text)
                        if rows_t:
                            rows    = rows_t
                            headers = headers_t
                            summary = "Extracted via pytesseract OCR"
                            logger.info(f"[Extractor] P4 pytesseract: {len(rows)} rows")
                except ImportError:
                    pass
                except Exception as te:
                    logger.debug(f"[Extractor] pytesseract: {te}")

            # ── Store result ──────────────────────────────────────────────
            if rows:
                result["extracted_data"]    = rows
                result["extracted_columns"] = headers
                result["extracted_text"]    = (
                    f"Image: {fp.name} | {summary} | "
                    f"{len(rows)} rows, columns: {headers}"
                )
                logger.info(f"[Extractor] Image done: {len(rows)} rows, cols={headers}")
            else:
                # Keep raw text so chat_agent can retry with Gemini
                result["extracted_text"]    = ocr_text or raw1 or ""
                result["extracted_data"]    = []
                result["extracted_columns"] = []
                logger.warning("[Extractor] Image: all passes got 0 rows")

        except Exception as e:
            from app.agents.quota_helper import is_quota_error
            logger.error(f"[Extractor] Image failed: {e}", exc_info=True)
            result["extracted_text"] = f"Image extraction error: {e}"
            if is_quota_error(e):
                raise

    # ── Word ──────────────────────────────────────────────────────────────────

    async def _extract_word(self, fp: Path, result: dict):
        try:
            from docx import Document
            doc        = Document(str(fp))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text_parts = ["=== DOCUMENT TEXT ==="] + paragraphs
            table_rows, all_cols = [], []
            for table in doc.tables:
                if not table.rows:
                    continue
                headers = [c.text.strip() for c in table.rows[0].cells]
                if not all_cols:
                    all_cols = headers
                for row in table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        table_rows.append(dict(zip(headers, cells)))
            result["extracted_text"]    = "\n".join(text_parts)
            result["extracted_data"]    = table_rows
            result["extracted_columns"] = all_cols
        except ImportError:
            try:
                import mammoth
                with open(fp, "rb") as f:
                    result["extracted_text"] = mammoth.extract_raw_text(f).value
            except Exception as e:
                result["extracted_text"] = f"Could not read Word document: {e}"

    # ── Plain text ────────────────────────────────────────────────────────────

    async def _extract_text(self, fp: Path, result: dict):
        text  = fp.read_text(encoding="utf-8", errors="ignore")
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        result["extracted_text"] = text
        for delim in [",", "\t", "|", ";"]:
            if lines and delim in lines[0]:
                headers = [h.strip() for h in lines[0].split(delim)]
                rows = []
                for line in lines[1:]:
                    vals = [v.strip() for v in line.split(delim)]
                    if len(vals) == len(headers):
                        rows.append({h: _coerce(v) for h, v in zip(headers, vals)})
                if rows:
                    result["extracted_data"]    = rows
                    result["extracted_columns"] = headers
                    return
        parsed = _parse_text_as_table(text)
        if parsed["rows"]:
            result["extracted_data"]    = parsed["rows"]
            result["extracted_columns"] = parsed["headers"]

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _build_summary(self, result: dict) -> str:
        parts = [f"File: {result['file_name']}"]
        if result["extracted_columns"]:
            parts.append(f"Columns: {result['extracted_columns']}")
        if result["row_count"]:
            parts.append(f"Rows: {result['row_count']}")
        elif result["extracted_text"]:
            parts.append(f"Text: {len(result['extracted_text'])} chars")
        return " | ".join(parts)

    def _empty_result(self, message: str) -> dict:
        return {
            "file_name": "", "file_type": "", "extracted_text": message,
            "extracted_data": [], "extracted_columns": [], "row_count": 0,
            "file_summary": message, "success": False,
        }


# ── Row merge helper ──────────────────────────────────────────────────────────

def _merge_rows(table_rows: list, gemini_rows: list, headers: list) -> list:
    """
    Merge two row lists, preferring table_rows and adding any rows from
    gemini_rows that aren't already in table_rows (matched by first column value).
    """
    if not headers:
        return table_rows or gemini_rows

    first_col = headers[0]
    seen_keys = {str(r.get(first_col, "")).strip().lower() for r in table_rows}

    merged = list(table_rows)
    for row in gemini_rows:
        key = str(row.get(first_col, "")).strip().lower()
        if key and key not in seen_keys:
            merged.append(row)
            seen_keys.add(key)

    return merged


# ── Gemini text→structured rows ───────────────────────────────────────────────

async def _gemini_extract_structured(text: str) -> list[dict]:
    """Call Gemini to extract rows from plain text. Quota errors re-raised."""
    try:
        from app.agents.quota_helper    import is_quota_error
        from app.services.gemini_service import gemini_service
        from app.agents.rate_limiter    import call_with_retry

        prompt = (
            "Extract ALL structured data rows from the text below.\n\n"
            "RULES:\n"
            "1. Use ONLY values present in the text. Never invent.\n"
            "2. Detect column names from headers in the text.\n"
            "3. Include EVERY data row — do not truncate.\n"
            "4. Numbers must be actual numbers.\n\n"
            f"TEXT:\n{text[:50000]}\n\n"
            "Return ONLY a valid JSON array:\n"
            '[{"Col1": "val", "Col2": 123}, ...]\n\n'
            "If no structured data: return []"
        )
        result = await call_with_retry(gemini_service.analyze_json, prompt)
        if isinstance(result, list) and result and isinstance(result[0], dict):
            logger.info(f"[Extractor] Gemini text: {len(result)} rows")
            return result
        return []
    except Exception as e:
        from app.agents.quota_helper import is_quota_error
        if is_quota_error(e):
            raise
        logger.warning(f"[Extractor] Gemini text extract failed: {e}")
        return []


# ── Vision call with retry ────────────────────────────────────────────────────

async def _vision_call_with_retry(model, parts, max_retries=3):
    """Run model.generate_content() in a thread with RPM auto-retry."""
    from app.agents.rate_limiter import _is_rpm_error, _is_daily_error, _get_retry_delay
    loop = asyncio.get_event_loop()
    last_exc = None
    for attempt in range(1, max_retries + 1):
        try:
            return await loop.run_in_executor(None, lambda: model.generate_content(parts))
        except Exception as e:
            if _is_daily_error(e):
                raise
            if _is_rpm_error(e):
                wait = _get_retry_delay(e) + 2
                logger.warning(
                    f"[Vision] RPM hit attempt {attempt}/{max_retries} "
                    f"— waiting {wait:.0f}s"
                )
                await asyncio.sleep(wait)
                last_exc = e
                continue
            raise
    raise last_exc


# ── Vision response parser — 5 layers ─────────────────────────────────────────

def _parse_vision_response(raw: str) -> tuple[list, list, str]:
    if not raw:
        return [], [], "Empty response"

    clean = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`").strip()

    # 1. JSON
    obj = None
    try:
        obj = json.loads(clean)
    except json.JSONDecodeError:
        m = re.search(r'\{.*\}', clean, re.DOTALL)
        if m:
            try:
                obj = json.loads(m.group())
            except Exception:
                pass

    if isinstance(obj, dict):
        headers  = obj.get("headers") or []
        rows_raw = obj.get("rows")    or []
        summary  = obj.get("summary", "")
        if headers and rows_raw:
            rows = []
            for row in rows_raw:
                if not isinstance(row, (list, tuple)):
                    continue
                padded   = list(row) + [""] * max(0, len(headers) - len(row))
                row_dict = {}
                for h, v in zip(headers, padded[:len(headers)]):
                    row_dict[h] = v if isinstance(v, (int, float)) else _coerce(str(v))
                rows.append(row_dict)
            if rows:
                logger.info(f"[Vision] JSON: {len(rows)} rows")
                return rows, headers, summary

    # 2. CSV / delimiter
    lines = [l.strip() for l in clean.splitlines() if l.strip()]
    for delim in [",", "\t", ";"]:
        delim_lines = [l for l in lines if delim in l]
        if len(delim_lines) < 2:
            continue
        col_counts = [l.count(delim) for l in delim_lines]
        if max(col_counts) < 1 or (max(col_counts) - min(col_counts)) > 2:
            continue
        parts0 = [p.strip() for p in delim_lines[0].split(delim)]
        is_hdr = all(not re.match(r'^-?[\d.]+$', p.strip()) for p in parts0 if p.strip())
        if is_hdr and len(delim_lines) >= 3:
            headers    = [p.strip().strip('"\'') for p in parts0 if p.strip()]
            data_lines = delim_lines[1:]
        else:
            headers    = [f"Col{i+1}" for i in range(len(parts0))]
            data_lines = delim_lines
        rows = []
        for line in data_lines:
            vals = [v.strip().strip('"\'') for v in line.split(delim)]
            if len(vals) < 2:
                continue
            if vals and re.match(r'^(total|summary|grand total)$', vals[0], re.I):
                continue
            while len(vals) < len(headers):
                vals.append("")
            row = {}
            for h, v in zip(headers, vals[:len(headers)]):
                v_c = re.sub(r'[$£€¥,]', '', v).strip()
                row[h] = _coerce(v_c) if v_c != v else _coerce(v)
            rows.append(row)
        if rows:
            logger.info(f"[Vision] CSV({delim!r}): {len(rows)} rows")
            return rows, headers, "Extracted from delimited text"

    # 3. Pipe / markdown table
    pipe_lines = [l.strip() for l in clean.splitlines() if "|" in l and l.strip()]
    if len(pipe_lines) >= 2:
        headers = [h.strip() for h in pipe_lines[0].split("|") if h.strip()]
        rows    = []
        for line in pipe_lines[1:]:
            if re.match(r'^[\|\s\-:]+$', line):
                continue
            vals = [v.strip() for v in line.split("|") if v.strip() != ""]
            if vals and len(vals) >= 2:
                while len(vals) < len(headers):
                    vals.append("")
                rows.append({h: _coerce(v) for h, v in zip(headers, vals[:len(headers)])})
        if rows:
            logger.info(f"[Vision] Pipe: {len(rows)} rows")
            return rows, headers, "Table"

    # 4. Name + number lines
    parsed = _parse_text_as_table(clean)
    if parsed["rows"]:
        logger.info(f"[Vision] Lines: {len(parsed['rows'])} rows")
        return parsed["rows"], parsed["headers"], "Lines"

    return [], [], "No data"


def _parse_text_as_table(text: str) -> dict:
    lines = [re.sub(r'\s+', ' ', l).strip() for l in text.splitlines() if l.strip()]
    name_score_rows = []
    for line in lines:
        line = line.rstrip(',').strip()
        m = re.match(r'^([A-Za-z][A-Za-z\s\-\.]{1,40}?)\s+([\d][\d\.\s]*)$', line)
        if m:
            name = m.group(1).strip()
            nums = re.findall(r'[\d]+(?:\.\d+)?', m.group(2))
            if nums:
                name_score_rows.append((name, nums))
    if len(name_score_rows) < 2:
        return {"headers": [], "rows": []}
    max_nums = max(len(r[1]) for r in name_score_rows)
    headers  = ["Name"] + (["Score"] if max_nums == 1 else [f"Score {i+1}" for i in range(max_nums)])
    rows = []
    for name, nums in name_score_rows:
        row = {"Name": name}
        for i, h in enumerate(headers[1:]):
            v = nums[i] if i < len(nums) else "0"
            row[h] = float(v) if '.' in v else int(v)
        rows.append(row)
    return {"headers": headers, "rows": rows}


def _coerce(val: str):
    if not isinstance(val, str):
        return val
    v = val.strip().replace(",", "")
    if v.lstrip("-").isdigit():
        return int(v)
    try:
        return float(v)
    except (ValueError, TypeError):
        return val


# Singleton
file_extractor = FileExtractor()
